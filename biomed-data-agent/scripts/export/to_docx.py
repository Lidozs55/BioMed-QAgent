"""Word 报告生成脚本（BioMed QAgent Stage 6，可选）。

输入：cleaned DataRecord JSON + 可选 lineage.json + 可选分析结果目录
输出：report.docx（Word 格式，中文）

复用 to_report.py 的 build_report() 构造报告结构，再用 python-docx 渲染为 Word。
python-docx 不可用时输出错误并建议安装该依赖或改用调度器自带的文档处理能力。

接口：
    python scripts/export/to_docx.py --input cleaned.json \
        --lineage lineage.json --analysis-dir results/ \
        --out report.docx --task-id T1

成功输出：{"status":"ok","output":"report.docx","rows":N}
失败输出：{"status":"error","message":"..."}
"""
import os
import sys

_HERE = os.path.dirname(os.path.abspath(__file__))
if _HERE not in sys.path:
    sys.path.insert(0, _HERE)
from _base import emit_error, emit_ok, load_lineage, load_records, log_stderr  # noqa: E402
from to_report import build_report  # noqa: E402


def render_docx(sections, output_path):
    """用 python-docx 把 sections 渲染为 Word 文档。"""
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    # 默认正文字体（中文兼容）
    try:
        doc.styles["Normal"].font.name = "Microsoft YaHei"
        doc.styles["Normal"].font.size = Pt(10.5)
    except Exception:
        pass
    for sec in sections:
        level = sec["level"]
        title = sec["title"]
        # 标题样式：level 1 → Heading 1, level 2 → Heading 2
        try:
            doc.add_heading(title, level=min(level, 3))
        except Exception:
            doc.add_paragraph(title)
        for b in sec["blocks"]:
            btype = b.get("type")
            if btype == "para":
                doc.add_paragraph(b.get("text", ""))
            elif btype == "bullets":
                for it in b.get("items", []):
                    doc.add_paragraph(it, style="List Bullet")
            elif btype == "table":
                headers = b.get("headers", [])
                rows = b.get("rows", [])
                tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
                tbl.style = "Light Grid Accent 1"
                for i, h in enumerate(headers):
                    tbl.rows[0].cells[i].text = str(h)
                for ri, row in enumerate(rows, start=1):
                    for ci, val in enumerate(row):
                        if ci < len(headers):
                            tbl.rows[ri].cells[ci].text = str(val)
    doc.save(output_path)


def main():
    import argparse
    parser = argparse.ArgumentParser(prog="to_docx", description="生成生物医学数据整合 Word 报告")
    parser.add_argument("--input", required=True, help="输入 cleaned DataRecord JSON")
    parser.add_argument("--lineage", default=None, help="lineage.json 路径（可选）")
    parser.add_argument("--analysis-dir", default=None, help="分析结果目录（可选）")
    parser.add_argument("--out", required=True, help="输出 report.docx 路径")
    parser.add_argument("--task-id", default="", help="任务 ID")
    args = parser.parse_args()
    try:
        try:
            import docx  # noqa: F401
        except ImportError:
            emit_error("python-docx 未安装。建议：(1) 沙箱中执行 pip install python-docx；"
                       "或 (2) 改用调度器自带的文档处理能力生成 Word 报告。")
        records = load_records(args.input)
        log_stderr(f"加载 {len(records)} 条记录")
        sections = build_report(records, load_lineage(args.lineage), args.task_id,
                                args.analysis_dir, args.out, input_path=args.input)
        os.makedirs(os.path.dirname(os.path.abspath(args.out)) or ".", exist_ok=True)
        render_docx(sections, args.out)
        log_stderr(f"Word 报告已写入 {args.out}（{len(sections)} 节）")
        emit_ok(args.out, rows=len(records), sections=len(sections))
    except Exception as e:
        emit_error(f"Word 报告生成失败: {e}")


if __name__ == "__main__":
    main()
